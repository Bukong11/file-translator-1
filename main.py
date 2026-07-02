from __future__ import annotations

import json
import math
import os
import shutil
import urllib.error
import urllib.request
import uuid
from pathlib import Path

import fitz
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from openpyxl import load_workbook
from pypdf import PdfReader


BASE_DIR = Path(__file__).resolve().parent
FRONTEND_DIR = BASE_DIR.parent / "frontend"
WORK_DIR = BASE_DIR / "storage"
WORK_DIR.mkdir(exist_ok=True)

SUPPORTED_SUFFIXES = {".docx", ".xlsx", ".pdf"}
OPENAI_API_URL = "https://api.openai.com/v1/responses"
DEFAULT_OPENAI_MODEL = "gpt-5.5"
LANGUAGE_NAMES = {
    "zh": "Chinese",
    "en": "English",
    "ms": "Malay",
    "ja": "Japanese",
    "ko": "Korean",
    "de": "German",
    "fr": "French",
}
PDF_FONT_FILE = next(
    (
        path
        for path in (
            Path(r"C:\Windows\Fonts\msyh.ttc"),
            Path(r"C:\Windows\Fonts\simhei.ttf"),
            Path(r"C:\Windows\Fonts\arial.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
        )
        if path.exists()
    ),
    None,
)

app = FastAPI(title="File Translator Backend")

allowed_origins = [
    origin.strip()
    for origin in os.getenv("ALLOWED_ORIGINS", "*").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)


def translate_with_openai(text: str, target_language: str, domain: str, terms: str) -> str:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    model = os.getenv("OPENAI_MODEL", DEFAULT_OPENAI_MODEL).strip()

    if not api_key:
        return ""

    prompt = (
        "Translate the text below. Keep the meaning accurate and professional. "
        "Return only the translated text. Do not add explanations, prefixes, quotes, or markdown.\n\n"
        f"Target language: {LANGUAGE_NAMES.get(target_language, target_language)}\n"
        f"Professional domain: {domain}\n"
        f"Terminology requirements: {terms or 'None'}\n\n"
        f"Text:\n{text}"
    )
    payload = {
        "model": model,
        "input": prompt,
    }
    request = urllib.request.Request(
        OPENAI_API_URL,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=60) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise HTTPException(status_code=502, detail=f"OpenAI API error: {detail}") from exc
    except urllib.error.URLError as exc:
        raise HTTPException(status_code=502, detail=f"Could not reach OpenAI API: {exc.reason}") from exc

    if data.get("output_text"):
        return data["output_text"].strip()

    pieces: list[str] = []
    for item in data.get("output", []):
        for content in item.get("content", []):
            if content.get("type") in {"output_text", "text"} and content.get("text"):
                pieces.append(content["text"])

    return "\n".join(pieces).strip()


def parse_json_array(text: str) -> list[str] | None:
    clean = text.strip()
    if clean.startswith("```"):
        clean = clean.strip("`").strip()
        if clean.lower().startswith("json"):
            clean = clean[4:].strip()

    try:
        data = json.loads(clean)
    except json.JSONDecodeError:
        start = clean.find("[")
        end = clean.rfind("]")
        if start < 0 or end <= start:
            return None
        try:
            data = json.loads(clean[start : end + 1])
        except json.JSONDecodeError:
            return None

    if not isinstance(data, list):
        return None
    return [str(item) for item in data]


def translate_many_with_openai(texts: list[str], target_language: str, domain: str, terms: str) -> list[str]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    model = os.getenv("OPENAI_MODEL", DEFAULT_OPENAI_MODEL).strip()
    if not api_key:
        label = LANGUAGE_NAMES.get(target_language, target_language)
        return [f"[{label} translation placeholder] {text}" for text in texts]

    prompt = (
        "Translate each item in the JSON array below. Keep technical terms accurate and professional. "
        "Return only a valid JSON array of strings, in the same order and with the same number of items. "
        "Do not add explanations, markdown, object keys, or extra text.\n\n"
        f"Target language: {LANGUAGE_NAMES.get(target_language, target_language)}\n"
        f"Professional domain: {domain}\n"
        f"Terminology requirements: {terms or 'None'}\n\n"
        f"JSON array:\n{json.dumps(texts, ensure_ascii=False)}"
    )
    payload = {"model": model, "input": prompt}
    request = urllib.request.Request(
        OPENAI_API_URL,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=90) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise HTTPException(status_code=502, detail=f"OpenAI API error: {detail}") from exc
    except urllib.error.URLError as exc:
        raise HTTPException(status_code=502, detail=f"Could not reach OpenAI API: {exc.reason}") from exc

    output_text = data.get("output_text", "")
    if not output_text:
        pieces: list[str] = []
        for item in data.get("output", []):
            for content in item.get("content", []):
                if content.get("type") in {"output_text", "text"} and content.get("text"):
                    pieces.append(content["text"])
        output_text = "\n".join(pieces).strip()

    translated = parse_json_array(output_text)
    if translated and len(translated) == len(texts):
        return translated

    return [translate_text(text, target_language, domain, terms) for text in texts]


def chunk_texts(texts: list[str], max_items: int = 30, max_chars: int = 5000) -> list[list[str]]:
    chunks: list[list[str]] = []
    current: list[str] = []
    current_chars = 0

    for text in texts:
        text_chars = len(text)
        if current and (len(current) >= max_items or current_chars + text_chars > max_chars):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(text)
        current_chars += text_chars

    if current:
        chunks.append(current)
    return chunks


def translate_text(text: str, target_language: str, domain: str, terms: str) -> str:
    clean = text.strip()
    if not clean:
        return text

    translated = translate_with_openai(text, target_language, domain, terms)
    if translated:
        return translated

    label = LANGUAGE_NAMES.get(target_language, target_language)
    return f"[{label} translation placeholder]"


def translate_docx(source: Path, target: Path, target_language: str, domain: str, terms: str) -> None:
    document = Document(source)

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            for run in paragraph.runs:
                if run.text.strip():
                    run.text = translate_text(run.text, target_language, domain, terms)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        for run in paragraph.runs:
                            if run.text.strip():
                                run.text = translate_text(run.text, target_language, domain, terms)

    document.save(target)


def translate_xlsx(source: Path, target: Path, target_language: str, domain: str, terms: str) -> None:
    workbook = load_workbook(source)
    cells_to_translate = []
    texts_to_translate: list[str] = []

    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if not should_translate_spreadsheet_cell(cell.value):
                    continue
                cells_to_translate.append(cell)
                texts_to_translate.append(str(cell.value))

    translated_texts: list[str] = []
    for chunk in chunk_texts(texts_to_translate):
        translated_texts.extend(translate_many_with_openai(chunk, target_language, domain, terms))

    for cell, translated in zip(cells_to_translate, translated_texts):
        cell.value = translated

    workbook.save(target)


def should_translate_spreadsheet_cell(value) -> bool:
    if not isinstance(value, str):
        return False

    clean = value.strip()
    if not clean or clean.startswith("="):
        return False

    letters = sum(ch.isalpha() for ch in clean)
    cjk = sum("\u4e00" <= ch <= "\u9fff" for ch in clean)
    digits = sum(ch.isdigit() for ch in clean)
    if letters + cjk == 0:
        return False
    if digits and digits / max(len(clean), 1) > 0.65:
        return False

    return True


def translate_pdf_to_docx(source: Path, target: Path, target_language: str, domain: str, terms: str) -> None:
    reader = PdfReader(str(source))
    document = Document()
    document.add_heading(source.stem, level=1)

    extracted_any_text = False
    for page_number, page in enumerate(reader.pages, start=1):
      text = page.extract_text() or ""
      clean = text.strip()
      if not clean:
          continue

      extracted_any_text = True
      document.add_heading(f"Page {page_number}", level=2)
      translated = translate_text(clean, target_language, domain, terms)
      for paragraph in translated.splitlines():
          if paragraph.strip():
              document.add_paragraph(paragraph.strip())

    if not extracted_any_text:
        document.add_paragraph(
            "No selectable text was found in this PDF. It may be a scanned image PDF and needs OCR before translation."
        )

    document.save(target)


def should_translate_pdf_text(text: str) -> bool:
    clean = " ".join(text.split())
    if len(clean) < 3:
        return False

    letters = sum(ch.isalpha() for ch in clean)
    cjk = sum("\u4e00" <= ch <= "\u9fff" for ch in clean)
    digits = sum(ch.isdigit() for ch in clean)

    if letters + cjk == 0:
        return False

    if digits and digits / max(len(clean), 1) > 0.65:
        return False

    protected_words = {
        "SUS430",
        "A3",
        "M5",
        "R3",
        "GB/T",
        "Alpha ESS",
    }
    if clean in protected_words:
        return False

    return True


def is_protected_drawing_table_area(page: fitz.Page, rect: fitz.Rect) -> bool:
    # Engineering drawings usually keep BOM, title block, revision table, item,
    # material, date, and signature fields in the lower band. These small cells
    # often cannot fit translated text, so preserve them exactly.
    return rect.y0 >= page.rect.height * 0.70


def line_rotation(line: dict, rect: fitz.Rect) -> int:
    if rect.width >= rect.height * 1.4:
        return 0

    direction = line.get("dir") or (1, 0)
    angle = math.degrees(math.atan2(direction[1], direction[0]))
    angle = (angle + 360) % 360
    nearest = round(angle / 90) * 90
    rotation = nearest % 360

    if rotation == 180:
        return 0
    if rect.height >= rect.width * 1.4 and rotation in {90, 270}:
        return rotation
    return 0


def fitted_font_size(rect: fitz.Rect, text: str, rotation: int = 0) -> float:
    available_height = rect.width if rotation in {90, 270} else rect.height
    height_based = max(3.5, min(8.0, available_height * 0.62))
    if len(text) > 120:
        return max(3.2, min(height_based, 5.0))
    if len(text) > 60:
        return max(3.4, min(height_based, 5.8))
    return height_based


def vertical_overlap(a: fitz.Rect, b: fitz.Rect) -> float:
    overlap = max(0.0, min(a.y1, b.y1) - max(a.y0, b.y0))
    return overlap / max(1.0, min(a.height, b.height))


def horizontal_safe_right_limit(page: fitz.Page, rect: fitz.Rect) -> float:
    search_band = fitz.Rect(rect.x0, rect.y0 - 4, page.rect.width - 12, rect.y1 + 4)
    right_limit = page.rect.width - 12

    for drawing in page.get_drawings():
        drawing_rect = drawing.get("rect")
        if not drawing_rect:
            continue
        obstacle = fitz.Rect(drawing_rect)
        if obstacle.x0 <= rect.x1 + 8:
            continue
        if vertical_overlap(search_band, obstacle) >= 0.15:
            right_limit = min(right_limit, obstacle.x0 - 8)

    for word in page.get_text("words"):
        obstacle = fitz.Rect(word[:4])
        if obstacle.x0 <= rect.x1 + 8:
            continue
        if vertical_overlap(search_band, obstacle) >= 0.40:
            right_limit = min(right_limit, obstacle.x0 - 8)

    return max(rect.x1, right_limit)


def candidate_text_rects(page: fitz.Page, rect: fitz.Rect, rotation: int) -> list[fitz.Rect]:
    if rotation in {90, 270}:
        return [rect]

    right_limit = horizontal_safe_right_limit(page, rect)
    line_height = max(4.0, rect.height + 2.0)
    return [
        rect,
        fitz.Rect(rect.x0, rect.y0 - 1, right_limit, rect.y0 - 1 + line_height),
        fitz.Rect(rect.x0, rect.y0 - 2, right_limit, rect.y0 - 2 + line_height + 1),
    ]


def choose_text_placement(page: fitz.Page, rect: fitz.Rect, text: str, rotation: int) -> tuple[fitz.Rect, float] | None:
    fontname = "msyh" if PDF_FONT_FILE else "helv"
    fontfile = str(PDF_FONT_FILE) if PDF_FONT_FILE else None
    base_size = fitted_font_size(rect, text, rotation)

    if rotation not in {90, 270}:
        candidate = candidate_text_rects(page, rect, rotation)[-1]
        text_width_at_one = max(1.0, fitz.get_text_length(text, fontname="helv", fontsize=1.0))
        font_size = min(base_size, candidate.width / text_width_at_one * 0.96)
        return candidate, max(1.35, font_size)

    for candidate in candidate_text_rects(page, rect, rotation):
        for font_size in (base_size, base_size * 0.9, base_size * 0.78, base_size * 0.64, 2.8, 2.4, 2.0, 1.7):
            scratch = fitz.open()
            scratch_page = scratch.new_page(width=page.rect.width, height=page.rect.height)
            result = scratch_page.insert_textbox(
                candidate,
                text,
                fontsize=font_size,
                fontname=fontname,
                fontfile=fontfile,
                color=(0, 0, 0),
                align=fitz.TEXT_ALIGN_LEFT,
                rotate=rotation,
                overlay=True,
            )
            scratch.close()
            if result >= 0:
                return candidate, font_size

    return None


def fallback_text_placement(page: fitz.Page, rect: fitz.Rect, rotation: int) -> tuple[fitz.Rect, float]:
    if rotation in {90, 270}:
        return rect, 2.2

    right_limit = horizontal_safe_right_limit(page, rect)
    line_height = max(4.0, rect.height + 2.0)
    fallback_rect = fitz.Rect(rect.x0, rect.y0 - 2, right_limit, rect.y0 - 2 + line_height)
    return fallback_rect, 1.6


def insert_translated_text(page: fitz.Page, rect: fitz.Rect, text: str, rotation: int, font_size: float) -> None:
    if rotation not in {90, 270}:
        page.insert_text(
            fitz.Point(rect.x0, rect.y1 - 1.0),
            text,
            fontsize=font_size,
            fontname="helv",
            color=(0, 0, 0),
            overlay=True,
        )
        return

    page.insert_textbox(
        rect,
        text,
        fontsize=font_size,
        fontname="msyh" if PDF_FONT_FILE else "helv",
        fontfile=str(PDF_FONT_FILE) if PDF_FONT_FILE else None,
        color=(0, 0, 0),
        align=fitz.TEXT_ALIGN_LEFT,
        rotate=rotation,
        overlay=True,
    )


def translate_pdf_overlay(source: Path, target: Path, target_language: str, domain: str, terms: str) -> None:
    document = fitz.open(source)

    for page in document:
        page_dict = page.get_text("dict")
        page_entries = []
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            for line in block.get("lines", []):
                text = " ".join(span.get("text", "") for span in line.get("spans", []))
                clean = " ".join(text.split())
                if not should_translate_pdf_text(clean):
                    continue

                rect = fitz.Rect(line.get("bbox"))
                rotation = line_rotation(line, rect)
                min_width = 5 if rotation in {90, 270} else 8
                min_height = 8 if rotation in {90, 270} else 5
                if rect.width < min_width or rect.height < min_height:
                    continue

                translated = translate_text(clean, target_language, domain, terms)
                if not translated or translated == clean:
                    continue

                placement = choose_text_placement(page, rect, translated, rotation)
                if not placement:
                    placement = fallback_text_placement(page, rect, rotation)

                page_entries.append((rect, placement[0], translated, rotation, placement[1]))

        for rect, _, _, _, _ in page_entries:
            cover = rect + (-0.8, -0.6, 0.8, 0.6)
            page.draw_rect(cover, color=None, fill=(1, 1, 1), overlay=True)

        for _, placement_rect, translated, rotation, font_size in page_entries:
            insert_translated_text(page, placement_rect, translated, rotation, font_size)

    document.save(target, garbage=4, deflate=True)
    document.close()


@app.get("/health")
def health() -> dict[str, str]:
    mode = "openai" if os.getenv("OPENAI_API_KEY", "").strip() else "placeholder"
    return {"status": "ok", "translation_mode": mode}


@app.get("/", response_model=None)
def root():
    index_file = FRONTEND_DIR / "index.html"
    if index_file.exists():
        return FileResponse(index_file)
    return {"name": "File Translator Backend", "status": "running", "health": "/health"}


@app.get("/config.js")
def frontend_config() -> Response:
    script = 'window.FILE_TRANSLATOR_API_BASE_URL = window.location.origin;\n'
    return Response(content=script, media_type="application/javascript")


@app.post("/translate")
async def translate_file(
    file: UploadFile = File(...),
    target_language: str = Form("zh"),
    domain: str = Form("business"),
    terms: str = Form(""),
    pdf_mode: str = Form("overlay"),
) -> FileResponse:
    original_name = Path(file.filename or "uploaded").name
    suffix = Path(original_name).suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        raise HTTPException(status_code=400, detail="This demo backend currently supports .docx, .xlsx, and text-based .pdf files.")

    task_id = uuid.uuid4().hex
    task_dir = WORK_DIR / task_id
    task_dir.mkdir()

    source = task_dir / original_name
    output_suffix = ".pdf" if suffix == ".pdf" and pdf_mode == "overlay" else ".docx" if suffix == ".pdf" else suffix
    output_name = f"{Path(original_name).stem}-translated{output_suffix}"
    target = task_dir / output_name

    with source.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    if suffix == ".docx":
        translate_docx(source, target, target_language, domain, terms)
    elif suffix == ".xlsx":
        translate_xlsx(source, target, target_language, domain, terms)
    elif suffix == ".pdf":
        if pdf_mode == "overlay":
            translate_pdf_overlay(source, target, target_language, domain, terms)
        else:
            translate_pdf_to_docx(source, target, target_language, domain, terms)

    return FileResponse(
        target,
        filename=output_name,
        media_type="application/octet-stream",
    )
